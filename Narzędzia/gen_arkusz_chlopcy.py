# -*- coding: utf-8 -*-
"""
Generator arkusza E8 — Chlopcy z Placu Broni
Fragment: wyprawa do Ogrodu Botanicznego, Nemeczek wpada do stawu
Data: 2026-03-22
"""

from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_PATH = r"D:\Dydaktyka\ArkuszeAI\Chlopcy_PlacBroni_OgrodBotaniczny_Arkusz_E8_AI.docx"

doc = Document()

# ── Marginesy 1.27 cm ──
for section in doc.sections:
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

# ── Style domyslne ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(2)

# ── Funkcje pomocnicze ──
def add_paragraph(text, bold=False, italic=False, alignment=None, space_after=Pt(2)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    return p

def add_mixed_paragraph(parts, alignment=None, space_after=Pt(2)):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    for text, b, i in parts:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = b
        run.italic = i
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    return p

def add_answer_line(count=1, space_after=Pt(2)):
    for _ in range(count):
        add_paragraph('_' * 90, space_after=space_after)

def set_cell_font(cell, text, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, size=Pt(12)):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = size
    run.bold = bold
    run.italic = italic
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(0)

def set_cell_mixed(cell, parts, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    for text, b, i in parts:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = b
        run.italic = i
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(0)

def remove_table_borders(table):
    """Ustawia brak obramowania dla calej tabeli."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is not None:
        tblPr.remove(borders)

def set_table_borders(table):
    """Ustawia pelne obramowanie tabeli."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is not None:
        tblPr.remove(borders)
    borders_xml = (
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    from lxml import etree
    borders_elem = etree.fromstring(borders_xml)
    tblPr.append(borders_elem)


# ══════════════════════════════════════════════════════════════
# NAGLOWEK ARKUSZA
# ══════════════════════════════════════════════════════════════
add_paragraph('Karta pracy — egzamin osmoklasisty', bold=True,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))

# Autor i tytul
add_mixed_paragraph([
    ('Ferenc Molnar, ', True, False),
    ('Chlopcy z Placu Broni', True, True),
], alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))

add_paragraph('(wyprawa do Ogrodu Botanicznego)', italic=True,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))

add_paragraph('Przeczytaj fragment i wykonaj zadania.', bold=False,
              space_after=Pt(6))

# ══════════════════════════════════════════════════════════════
# FRAGMENT — PROZA, TNR 12pt, space_after=Pt(2)
# ══════════════════════════════════════════════════════════════

fragment_lines = [
    "Wreszcie byli na miejscu. Tu juz mogli sie wyprostowac: sitowie, trzciny i nadbrzezne krzaki byly tak wysokie, ze calkiem ich zakrywaly. Boka nie tracac zimnej krwi wydawal komendy.",
    "- Gdzies tu, w poblizu, powinna znajdowac sie lodka. Ja z Nemeczkiem pojde szukac lodzi w prawo, a ty, Czonakosz na lewo. Kto pierwszy znajdzie lodke, ten niech zaczeka przy niej.",
    "Natychmiast ruszyli na poszukiwanie. Ledwo zrobili kilka krokow, Boka dojrzal lodke wsrod sitowia.",
    "- Zaczekajmy - powiedzial.",
    "Czekali, az Czonakosz okrazy caly staw i nadejdzie z drugiej strony. Usiedli na brzegu i przez chwile patrzyli na gwiazdziste niebo. Potem zaczeli nadstawiac uszu z nadzieja, ze uslysza prowadzone na wyspie rozmowy. Nemeczek chcial sie popisac przed Boka.",
    "- Sluchaj - powiedzial - przyloze ucho do ziemi, to moze cos uslysze.",
    "- Daj spokoj - odpowiedzial Boka. - Nie ma co przykladac ucha do ziemi nad stawem. Ale jesli nachylimy sie nad woda, to bedziemy wszystko slyszec. Widzialem, jak rybacy na Dunaju w taki sposob porozumiewali sie ze soba z obu brzegow rzeki. Zwlaszcza wieczorem glos doskonale niesie sie po wodzie.",
    "Pochylili sie nad powierzchnia stawu, ale zadnego slowa nie uslyszeli wyraznie. Z wyspy dobiegaly tylko jakies sciszone glosy i szmery. Wreszcie nadszedl Czonakosz i zrozpaczony zamelddowal:",
    "- Nigdzie nie ma tej lodzi.",
    "- Nie martw sie, staruszku - pocieszyl go Nemeczek - znalezlismy ja.",
    "Ruszyli razem w strone lodki.",
    "- Wsiadamy?",
    "- Nie tutaj - postanowil Boka. - Naprzod musimy przeholowac lodz na przeciwlegly brzeg, zebysmy nie znalezli sie w poblizu tego mostu, jesli nas zauwaza. Musimy przeprawic sie na wyspe w miejscu najbardziej odleglym od mostu. Wtedy beda musieli nadrobic duzy kawalek, zeby nas dogonic.",
    "Ta przezornosc spodobala sie obu chlopcom. Swiadomosc, ze ich dowodca potrafi tak madrze wszystko przewidywac, dodala im otuchy. Boka odezwal sie:",
    "- Czy ktos z was ma linke?",
    "Czonakosz mial. Zreszta w kieszeni Czonakosza bylo wszystko. Nie ma takiego straganu na swiecie, na ktorym byloby tyle najrozmaitszych rzeczy, co w kieszeniach Czonakosza. A wiec mial on przy sobie scyzoryk, sznurek, kulki, mosiezna klamke, gwozdzie, klucze, szmaty, notes, srubokret i Bog jeden wie, co jeszcze. Czonakosz wydobyl linke, a Boka przywi\u0105zal ja do znajdującego sie na dziobie lodzi kolka. Po czym ostroznie i powoli zaczeli ciagnac lodz wzdluz brzegu, w kierunku przeciwleglej strony wyspy. W czasie holowania caly czas bacznie obserwowali wyspe. Kiedy wreszcie dobrn\u0119li do tego miejsca, z ktorego chcieli sie przeprawic, znow uslyszeli gwizd, taki sam jak poprzednio. Tym razem jednak juz sie nie przestraszyli. Wiedzieli bowiem, ze ten gwizd oznacza zmiane warty. Nie bali sie rowniez i z tego wzgledu, ze poczuli sie nagle w ogniu walki. Tak samo bywa z zolnierzami w czasie prawdziwej wojny. Poki nie zobacza wroga, boja sie byle krzaka. Kiedy jednak pierwsza kula swisnie im kolo ucha, nabieraja odwagi, ryzyko walki wciaga ich bez reszty i nie mysla o tym, ze ida po smierc.",
    "Chlopcy zaczeli wsiadac do lodki. Pierwszy wskoczyl Boka. Za nim Czonakosz. Nemeczek lekliwie dreptal po mulistym brzegu.",
    "- Chodzze, staruszku, chodzze - zachecal go Czonakosz.",
    "- Ide, staruszku, ide - odpowiedzial Nemeczek i w tym momencie poslizgnal sie, ze strachu chwycil trzcine i nie pisnawszy nawet wpadl do wody. Zanuzyl sie az po szyje. Nie krzyknal jednak i natychmiast sie podniosl. Bylo tam plytko. Woda sciekala z niego strugami, w dloni trzymal wciaz kurczowo cieniutka trzcine i wygladal bardzo smiesznie.",
    "Czonakosz nie mogl powstrzymac sie od smiechu.",
    "- Napiles sie, staruszku? - parsknal.",
    "- Nie napilem - odpowiedzial z przerazona mina blondynek i zablocony, przemokniety do suchej nitki, wgramolil sie do lodzi. Byl calkiem blady z wrazenia.",
    "- Nie myslalem, ze dzis jeszcze bede sie kapal - dodal cicho.",
    "Nie bylo chwili do stracenia. Boka i Czonakosz chwycili za wiosla i odepchneli lodz od brzegu. Ciezka lajba leniwie wyplynela na staw, marszczac gladka powierzchnie wody. Bezszelestnie zanurzali piora wiosel. Panowala tak wielka cisza, ze slychac bylo, jak siedzacemu na dziobie Nemeczkowi szczekaja z zimna zeby. Po kilku chwilach dziob lodzi dotknal brzegu wyspy. Chlopcy w pospiechu wyskoczyli na lad i natychmiast ukryli sie w krzakach.",
    "- No, wreszcie. Jestesmy - powiedzial Boka i ostroznie, cicho zaczal sie skradac wzdluz brzegu. Obaj chlopcy za nim.",
    "- Hej! - dowodca odwrocil sie nagle - nie mozemy przeciez zostawic lodzi bez opieki. Jesli odkryja lodz, to nie bedziemy mogli wrocic z wyspy. Ty, Czonakosz, zostaniesz przy czolnie. A gdyby ktos zauwazyl lodz, wsadz palce do ust i gwizdnij najglosniej, jak potrafisz. Wrocimy wtedy pedem, wskoczymy do lodzi, a ty ja odepchniesz od brzegu.",
    "Czonakosz smetnie wrocil do czolna, ale w skrytosci ducha cieszyl sie, ze moze nadarzy sie okazja do takiego gwizdniecia, ile sil w plucach..."
]

# Zamiast transliteracji — uzyje ORYGINALNEGO tekstu z pliku z normalizacja Ŝ->ż
# Wczytam fragment ponownie z oryginalnymi polskimi znakami

# Oryginalny tekst z pliku (z zamiana Ŝ na ż):
original_fragment = """Wreszcie byli na miejscu. Tu już mogli się wyprostować: sitowie, trzciny i nadbrzeżne krzaki były tak wysokie, że całkiem ich zakrywały. Boka nie tracąc zimnej krwi wydawał komendy.
- Gdzieś tu, w pobliżu, powinna znajdować się łódka. Ja z Nemeczkiem pójdę szukać łodzi w prawo, a ty, Czonakosz na lewo. Kto pierwszy znajdzie łódkę, ten niech zaczeka przy niej.
Natychmiast ruszyli na poszukiwanie. Ledwo zrobili kilka kroków, Boka dojrzał łódkę wśród sitowia.
- Zaczekajmy - powiedział.
Czekali, aż Czonakosz okrąży cały staw i nadejdzie z drugiej strony. Usiedli na brzegu i przez chwilę patrzyli na gwiaździste niebo. Potem zaczęli nadstawiać uszu z nadzieją, że usłyszą prowadzone na wyspie rozmowy. Nemeczek chciał się popisać przed Boką.
- Słuchaj - powiedział - przyłożę ucho do ziemi, to może coś usłyszę.
- Daj spokój - odpowiedział Boka. - Nie ma co przykładać ucha do ziemi nad stawem. Ale jeśli nachylimy się nad wodą, to będziemy wszystko słyszeć. Widziałem, jak rybacy na Dunaju w taki sposób porozumiewali się ze sobą z obu brzegów rzeki. Zwłaszcza wieczorem głos doskonale niesie się po wodzie.
Pochylili się nad powierzchnią stawu, ale żadnego słowa nie usłyszeli wyraźnie. Z wyspy dobiegały tylko jakieś ściszone głosy i szmery. Wreszcie nadszedł Czonakosz i zrozpaczony zameldował:
- Nigdzie nie ma tej łodzi.
- Nie martw się, staruszku - pocieszył go Nemeczek - znaleźliśmy ją.
Ruszyli razem w stronę łódki.
- Wsiadamy?
- Nie tutaj - postanowił Boka. - Naprzód musimy przeholować łódź na przeciwległy brzeg, żebyśmy nie znaleźli się w pobliżu tego mostu, jeśli nas zauważą. Musimy przeprawić się na wyspę w miejscu najbardziej odległym od mostu. Wtedy będą musieli nadrobić duży kawałek, żeby nas dogonić.
Ta przezorność spodobała się obu chłopcom. Świadomość, że ich dowódca potrafi tak mądrze wszystko przewidywać, dodała im otuchy. Boka odezwał się:
- Czy ktoś z was ma linkę?
Czonakosz miał. Zresztą w kieszeni Czonakosza było wszystko. Nie ma takiego straganu na świecie, na którym byłoby tyle najrozmaitszych rzeczy, co w kieszeniach Czonakosza. A więc miał on przy sobie scyzoryk, sznurek, kulki, mosiężną klamkę, gwoździe, klucze, szmaty, notes, śrubokręt i Bóg jeden wie, co jeszcze. Czonakosz wydobył linkę, a Boka przywiązał ją do znajdującego się na dziobie łodzi kółka. Po czym ostrożnie i powoli zaczęli ciągnąć łódź wzdłuż brzegu, w kierunku przeciwległej strony wyspy. W czasie holowania cały czas bacznie obserwowali wyspę. Kiedy wreszcie dobrnęli do tego miejsca, z którego chcieli się przeprawić, znów usłyszeli gwizd, taki sam jak poprzednio. Tym razem jednak już się nie przestraszyli. Wiedzieli bowiem, że ten gwizd oznacza zmianę warty. Nie bali się również i z tego względu, że poczuli się nagle w ogniu walki. Tak samo bywa z żołnierzami w czasie prawdziwej wojny. Póki nie zobaczą wroga, boją się byle krzaka. Kiedy jednak pierwsza kula świśnie im koło ucha, nabierają odwagi, ryzyko walki wciąga ich bez reszty i nie myślą o tym, że idą po śmierć.
Chłopcy zaczęli wsiadać do łódki. Pierwszy wskoczył Boka. Za nim Czonakosz. Nemeczek lękliwie dreptał po mulistym brzegu.
- Chodźże, staruszku, chodźże - zachęcał go Czonakosz.
- Idę, staruszku, idę - odpowiedział Nemeczek i w tym momencie poślizgnął się, ze strachu chwycił trzcinę i nie pisnąwszy nawet wpadł do wody. Zanurzył się aż po szyję. Nie krzyknął jednak i natychmiast się podniósł. Było tam płytko. Woda ściekała z niego strugami, w dłoni trzymał wciąż kurczowo cieniutką trzcinę i wyglądał bardzo śmiesznie.
Czonakosz nie mógł powstrzymać się od śmiechu.
- Napiłeś się, staruszku? - parsknął.
- Nie napiłem - odpowiedział z przerażoną miną blondynek i zabłocony, przemoczony do suchej nitki, wgramolił się do łodzi. Był całkiem blady z wrażenia.
- Nie myślałem, że dziś jeszcze będę się kąpał - dodał cicho.
Nie było chwili do stracenia. Boka i Czonakosz chwycili za wiosła i odepchnęli łódź od brzegu. Ciężka łajba leniwie wypłynęła na staw, marszcząc gładką powierzchnię wody. Bezszelestnie zanurzali pióra wioseł. Panowała tak wielka cisza, że słychać było, jak siedzącemu na dziobie Nemeczkowi szczękają z zimna zęby. Po kilku chwilach dziób łodzi dotknął brzegu wyspy. Chłopcy w pośpiechu wyskoczyli na ląd i natychmiast ukryli się w krzakach.
- No, wreszcie. Jesteśmy - powiedział Boka i ostrożnie, cicho zaczął się skradać wzdłuż brzegu. Obaj chłopcy za nim.
- Hej! - dowódca odwrócił się nagle - nie możemy przecież zostawić łodzi bez opieki. Jeśli odkryją łódź, to nie będziemy mogli wrócić z wyspy. Ty, Czonakosz, zostaniesz przy czółnie. A gdyby ktoś zauważył łódź, wsadź palce do ust i gwizdnij najgłośniej, jak potrafisz. Wrócimy wtedy pędem, wskoczymy do łodzi, a ty ją odepchniesz od brzegu.
Czonakosz smętnie wrócił do czółna, ale w skrytości ducha cieszył się, że może nadarzy się okazja do takiego gwizdnięcia, ile sił w płucach..."""

# Podziel na akapity
paragraphs = [p.strip() for p in original_fragment.split('\n') if p.strip()]

for para_text in paragraphs:
    p = doc.add_paragraph()
    run = p.add_run(para_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    # Interlinia single (proza)
    p.paragraph_format.line_spacing = Pt(14)

# ══════════════════════════════════════════════════════════════
# ZADANIA
# ══════════════════════════════════════════════════════════════

# Separator
add_paragraph('', space_after=Pt(6))

# ── ZADANIE 1 (P/F) ──
add_paragraph('Zadanie 1.', bold=True, space_after=Pt(2))
add_paragraph('Zdecyduj, czy podane zdania sa zgodne z trescia fragmentu. Zaznacz P (prawda) lub F (falsz).', space_after=Pt(4))

# Poprawna wersja z polskimi znakami
p1 = doc.add_paragraph()
run1 = p1.add_run('Zadanie 1.')
run1.font.name = 'Times New Roman'
run1.font.size = Pt(12)
run1.bold = True
p1.paragraph_format.space_after = Pt(2)

# Usun zduplikowane — zacznijmy od nowa
# Usuwam dwa ostatnie paragrafy (zduplikowane zadanie 1)
# Prostsze: nadpiszmy caly plik prawidlowo

# Zbudujmy od nowa
doc2 = Document()
doc = doc2

for section in doc.sections:
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(2)

# ── NAGLOWEK ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Karta pracy — egzamin osmoklasisty')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.bold = True
p.paragraph_format.space_after = Pt(6)

# Autor + tytul
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = p.add_run('Ferenc Molnar, ')
r1.font.name = 'Times New Roman'
r1.font.size = Pt(12)
r1.bold = True
r2 = p.add_run('Chlopcy z Placu Broni')
r2.font.name = 'Times New Roman'
r2.font.size = Pt(12)
r2.bold = True
r2.italic = True
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('(wyprawa do Ogrodu Botanicznego)')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.italic = True
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
run = p.add_run('Przeczytaj fragment i wykonaj zadania.')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(6)

# ── FRAGMENT ──
for para_text in paragraphs:
    p = doc.add_paragraph()
    run = p.add_run(para_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(14)

# Separator
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)

# ══════════════════════════════════
# ZADANIE 1 (P/F)
# ══════════════════════════════════
p = doc.add_paragraph()
run = p.add_run('Zadanie 1.')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.bold = True
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
run = p.add_run('Zdecyduj, czy podane zdania są zgodne z treścią fragmentu. Zaznacz ')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
r2 = p.add_run('P')
r2.font.name = 'Times New Roman'
r2.font.size = Pt(12)
r2.bold = True
r3 = p.add_run(' (prawda) lub ')
r3.font.name = 'Times New Roman'
r3.font.size = Pt(12)
r4 = p.add_run('F')
r4.font.name = 'Times New Roman'
r4.font.size = Pt(12)
r4.bold = True
r5 = p.add_run(' (fałsz).')
r5.font.name = 'Times New Roman'
r5.font.size = Pt(12)
p.paragraph_format.space_after = Pt(4)

# Tabela P/F
table1 = doc.add_table(rows=3, cols=3)
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(table1)

# Naglowek tabeli
set_cell_font(table1.rows[0].cells[0], 'Zdanie', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_font(table1.rows[0].cells[1], 'P', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_font(table1.rows[0].cells[2], 'F', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

# Wiersz 1
set_cell_font(table1.rows[1].cells[0],
    'Boka wiedział, że głos dobrze niesie się po wodzie, bo obserwował rybaków na Dunaju.')
set_cell_font(table1.rows[1].cells[1], '', alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_font(table1.rows[1].cells[2], '', alignment=WD_ALIGN_PARAGRAPH.CENTER)

# Wiersz 2
set_cell_font(table1.rows[2].cells[0],
    'Chłopcy dotarli na wyspę, przekraczając strzeżony most.')
set_cell_font(table1.rows[2].cells[1], '', alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_font(table1.rows[2].cells[2], '', alignment=WD_ALIGN_PARAGRAPH.CENTER)

# Szerokosc kolumn
from docx.shared import Cm as CmShare
table1.columns[0].width = Cm(14)
table1.columns[1].width = Cm(1)
table1.columns[2].width = Cm(1)
for row in table1.rows:
    row.cells[0].width = Cm(14)
    row.cells[1].width = Cm(1)
    row.cells[2].width = Cm(1)

# Separator
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)

# ══════════════════════════════════
# ZADANIE 2 (WW)
# ══════════════════════════════════
p = doc.add_paragraph()
run = p.add_run('Zadanie 2.')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.bold = True
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
run = p.add_run('Co Czonakosz miał w kieszeniach podczas wyprawy? Zaznacz właściwą odpowiedź.')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(2)

odpowiedzi_zad2 = [
    'A. Lornetkę teatralną i osiem tomahawków.',
    'B. Scyzoryk, sznurek, kulki, mosiężną klamkę, gwoździe, klucze, szmaty, notes, śrubokręt i inne.',
    'C. Linkę, notes i mapy Ogrodu Botanicznego.',
    'D. Trzy tomahawki i mosiężną klamkę.',
]

for odp in odpowiedzi_zad2:
    p = doc.add_paragraph()
    run = p.add_run(odp)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(2)

# Separator
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)

# ══════════════════════════════════
# ZADANIE 3 (SLOWNIK)
# ══════════════════════════════════
p = doc.add_paragraph()
run = p.add_run('Zadanie 3.')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.bold = True
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run('W zdaniu: „Czonakosz smętnie wrócił do czółna, ale ')
r.font.name = 'Times New Roman'
r.font.size = Pt(12)
r2 = p.add_run('w skrytości ducha')
r2.font.name = 'Times New Roman'
r2.font.size = Pt(12)
r2.bold = True
r3 = p.add_run(' cieszył się, że może nadarzy się okazja do takiego gwizdnięcia, ile sił w płucach" — wyjaśnij znaczenie podkreślonego związku wyrazowego ')
r3.font.name = 'Times New Roman'
r3.font.size = Pt(12)
r4 = p.add_run('w skrytości ducha')
r4.font.name = 'Times New Roman'
r4.font.size = Pt(12)
r4.bold = True
r5 = p.add_run('.')
r5.font.name = 'Times New Roman'
r5.font.size = Pt(12)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
run = p.add_run('_' * 90)
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(2)

# Separator
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)

# ══════════════════════════════════
# ZADANIE 4 (INTERPUNKCJA)
# ══════════════════════════════════
p = doc.add_paragraph()
run = p.add_run('Zadanie 4.')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.bold = True
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r = p.add_run('Zapisz, jaką funkcję pełnią przecinki w zdaniu: „A więc miał on przy sobie scyzoryk, sznurek, kulki, mosiężną klamkę, gwoździe, klucze, szmaty, notes, śrubokręt i Bóg jeden wie, co jeszcze". Odpowiedz pełnym zdaniem.')
r.font.name = 'Times New Roman'
r.font.size = Pt(12)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
run = p.add_run('_' * 90)
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(2)

# Separator
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)

# ══════════════════════════════════
# ZADANIE 5 (INNA_LEKTURA)
# ══════════════════════════════════
p = doc.add_paragraph()
run = p.add_run('Zadanie 5.')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.bold = True
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r1 = p.add_run('Spośród lektur obowiązkowych — innych niż ')
r1.font.name = 'Times New Roman'
r1.font.size = Pt(12)
r2 = p.add_run('Chłopcy z Placu Broni')
r2.font.name = 'Times New Roman'
r2.font.size = Pt(12)
r2.italic = True
r3 = p.add_run(' — wybierz tę, w której bohater wykazuje odwagę, działając mimo strachu lub trudnych okoliczności. Podaj tytuł lektury i imię bohatera. Uzasadnij wybór, przywołując konkretną sytuację z tekstu.')
r3.font.name = 'Times New Roman'
r3.font.size = Pt(12)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
run = p.add_run('Bohater i tytuł lektury: ' + '_' * 70)
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
run = p.add_run('Uzasadnienie:')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(2)

for _ in range(3):
    p = doc.add_paragraph()
    run = p.add_run('_' * 90)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(2)

# Separator
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)

# ══════════════════════════════════
# ZADANIE 6 (PP — narrator)
# ══════════════════════════════════
p = doc.add_paragraph()
run = p.add_run('Zadanie 6.')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.bold = True
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r1 = p.add_run('Chłopcy z Placu Broni')
r1.font.name = 'Times New Roman'
r1.font.size = Pt(12)
r1.italic = True
r2 = p.add_run(' to powieść — utwór epicki z narratorem.')
r2.font.name = 'Times New Roman'
r2.font.size = Pt(12)
p.paragraph_format.space_after = Pt(4)

# Podpunkt a)
p = doc.add_paragraph()
r = p.add_run('a) Określ typ narratora prowadzącego opowiadanie w tej powieści. Uzasadnij odpowiedź jednym przykładem z przytoczonego fragmentu.')
r.font.name = 'Times New Roman'
r.font.size = Pt(12)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
run = p.add_run('Typ narratora: ' + '_' * 40)
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
run = p.add_run('Uzasadnienie: ' + '_' * 76)
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(4)

# Podpunkt b)
p = doc.add_paragraph()
r = p.add_run('b) Podaj jedną funkcję narratora w tym fragmencie — co czytelnik dowiaduje się od narratora poza tym, co wynika z dialogów postaci? Odpowiedz pełnym zdaniem.')
r.font.name = 'Times New Roman'
r.font.size = Pt(12)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
run = p.add_run('_' * 90)
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(2)

# Separator
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)

# ══════════════════════════════════
# ZADANIE 7 (CHRONOLOGIA)
# ══════════════════════════════════
p = doc.add_paragraph()
run = p.add_run('Zadanie 7.')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.bold = True
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
r1 = p.add_run('Ułóż poniższe zdarzenia z powieści ')
r1.font.name = 'Times New Roman'
r1.font.size = Pt(12)
r2 = p.add_run('Chłopcy z Placu Broni')
r2.font.name = 'Times New Roman'
r2.font.size = Pt(12)
r2.italic = True
r3 = p.add_run(' w kolejności, w jakiej pojawiają się w całym utworze. Wpisz cyfry 1–5 w kratki (1 = najwcześniejsze, 5 = najpóźniejsze).')
r3.font.name = 'Times New Roman'
r3.font.size = Pt(12)
p.paragraph_format.space_after = Pt(4)

# Tabela chronologii
table7 = doc.add_table(rows=6, cols=3)
table7.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(table7)

# Naglowek
set_cell_font(table7.rows[0].cells[0], '', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_font(table7.rows[0].cells[1], 'Zdarzenie', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_font(table7.rows[0].cells[2], 'Kolejność', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

zdarzenia = [
    ('A.', 'Geréb zdradza kolegów i przechodzi na stronę czerwonych koszul.'),
    ('B.', 'Odbywa się decydująca bitwa o Plac Broni.'),
    ('C.', 'Nemeczek umiera na zapalenie płuc.'),
    ('D.', 'Chłopcy z Placu Broni dowiadują się, że ich teren ma zostać sprzedany.'),
    ('E.', 'Nemeczek zostaje jedynym szeregowcem w drużynie Boki.'),
]

for i, (lit, zd) in enumerate(zdarzenia, 1):
    set_cell_font(table7.rows[i].cells[0], lit, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_font(table7.rows[i].cells[1], zd)
    set_cell_font(table7.rows[i].cells[2], '', alignment=WD_ALIGN_PARAGRAPH.CENTER)

# Szerokosc kolumn
table7.columns[0].width = Cm(1)
table7.columns[1].width = Cm(13)
table7.columns[2].width = Cm(1.5)
for row in table7.rows:
    row.cells[0].width = Cm(1)
    row.cells[1].width = Cm(13)
    row.cells[2].width = Cm(1.5)

# ══════════════════════════════════
# ZAPIS
# ══════════════════════════════════
doc.save(OUTPUT_PATH)
print(f"Arkusz zapisany: {OUTPUT_PATH}")
