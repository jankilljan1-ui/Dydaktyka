# -*- coding: utf-8 -*-
"""
Generator arkusza E8 — Ch\u0142opcy z Placu Broni, Rozdz. II
Nemeczek na Placu, odkrycie Feriego Acza, kradzie\u017c chor\u0105giewki
id=8
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# === MARGINESY 1.27 cm ===
for section in doc.sections:
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.space_before = Pt(0)

LINE = '_' * 90


def add_paragraph(text, bold=False, italic=False, alignment=None, space_after=Pt(2)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    return p


def add_mixed_paragraph(parts, alignment=None, space_after=Pt(2)):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    return p


# === NAGLOWEK ===
add_paragraph(
    'Karta pracy \u2014 egzamin \u00f3smoklasisty',
    bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6)
)

add_mixed_paragraph([
    ('Ferenc Moln\u00e1r, ', False, False),
    ('Ch\u0142opcy z Placu Broni', False, True),
    (', rozdzia\u0142 II', False, False)
], alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))

add_paragraph(
    'Przeczytaj uwa\u017cnie poni\u017cszy fragment, a nast\u0119pnie wykonaj zadania.',
    italic=True, space_after=Pt(6)
)

# === WCZYTANIE ORYGINALNEGO FRAGMENTU Z PLIKU ===
import glob as _glob
_candidates = _glob.glob(r"D:\Dydaktyka\Lektury\Fragmenty\ch*placubroni_TEKST.txt")
if not _candidates:
    raise FileNotFoundError("Nie znaleziono pliku ch*placubroni_TEKST.txt")
source_path = _candidates[0]
print(f"Plik zrodlowy: {source_path}")

with open(source_path, 'r', encoding='utf-8') as f:
    all_lines = f.readlines()

# Znajdz poczatek i koniec fragmentu
start_idx = None
end_idx = None
for i, line in enumerate(all_lines):
    if 'Kilka minut po p' in line and 'trzeciej skrzypn' in line and 'furtka od ulicy Paw' in line:
        start_idx = i
    if start_idx is not None and 'chor\u0105giewka uszyta przez siostr' in line:
        end_idx = i
        break

if start_idx is None or end_idx is None:
    # Fallback: szukaj po uproszczonych fragmentach
    for i, line in enumerate(all_lines):
        if start_idx is None and 'Kilka minut po' in line and 'Nemeczek' in line:
            start_idx = i
        if start_idx is not None and 'uszyta przez siostr' in line:
            end_idx = i
            break

if start_idx is None or end_idx is None:
    raise ValueError(f"Nie znaleziono fragmentu! start={start_idx}, end={end_idx}")

print(f"Fragment: linie {start_idx+1}--{end_idx+1} (plik)")

# Filtruj linie
fragment_raw = []
in_footnote = False
for i in range(start_idx, end_idx + 1):
    line = all_lines[i].rstrip('\n').rstrip()
    # Pomin markery stron
    if line.startswith('--- STRONA'):
        continue
    # Puste linie
    if line.strip() == '':
        if not in_footnote:
            fragment_raw.append('')
        continue
    # Pomin przypis tlumacza (zaczyna sie od gwiazdki i zawiera slowo o kolorach)
    if line.strip().startswith('\u2217') and ('Czerwono' in line or 'czerwono' in line):
        in_footnote = True
        continue
    if in_footnote:
        # Kontynuacja przypisu — linie z nawiasami, "przypisy", "zielonych"
        if ('przypisy' in line.lower() or 'zielonych' in line.lower()
                or line.strip().startswith('(') or 'Kossutha' in line
                or 'barwami' in line or 'narodow' in line.lower()):
            continue
        else:
            in_footnote = False
    if not in_footnote:
        fragment_raw.append(line)

# Wstaw do dokumentu
for line in fragment_raw:
    if line.strip() == '':
        continue
    p = doc.add_paragraph()
    text = line
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

# Separator
add_paragraph('', space_after=Pt(6))

# === ZADANIE 1 (P/F) ===
add_mixed_paragraph([
    ('Zadanie 1. ', True, False),
    ('Zdecyduj, czy poni\u017csze zdania s\u0105 prawdziwe (P), czy fa\u0142szywe (F). W\u0142a\u015bciw\u0105 odpowied\u017a zaznacz w tabeli.', False, False)
], space_after=Pt(4))

# Tabela P/F — 2 zdania
table = doc.add_table(rows=3, cols=3)
table.style = 'Table Grid'

# Naglowki
cells_h = table.rows[0].cells
for ci, ht in enumerate(['', 'P', 'F']):
    cells_h[ci].text = ''
    p = cells_h[ci].paragraphs[0]
    run = p.add_run(ht)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if ci > 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

zdania_pf = [
    'Nemeczek jako pierwszy przyszed\u0142 na zebranie i starannie zamkn\u0105\u0142 za sob\u0105 furtk\u0119, zgodnie z regulaminem Placu.',
    'Wspinaj\u0105c si\u0119 na szczyt twierdzy, Nemeczek milcza\u0142 i niczym nie zdradza\u0142, \u017ce si\u0119 boi.'
]
for row_idx, zd in enumerate(zdania_pf, start=1):
    cells = table.rows[row_idx].cells
    cells[0].text = ''
    p = cells[0].paragraphs[0]
    run = p.add_run(zd)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    for ci in [1, 2]:
        cells[ci].text = ''
        cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Szerokosci kolumn
for row in table.rows:
    row.cells[0].width = Cm(14)
    row.cells[1].width = Cm(1.5)
    row.cells[2].width = Cm(1.5)

add_paragraph('', space_after=Pt(4))

# === ZADANIE 2 (WW) ===
add_mixed_paragraph([
    ('Zadanie 2. ', True, False),
    ('Zaznacz w\u0142a\u015bciw\u0105 odpowied\u017a.', False, False)
], space_after=Pt(2))

add_paragraph(
    'Dlaczego narrator sugeruje, \u017ce Nemeczek tak bardzo polubi\u0142 psa Hektora?',
    space_after=Pt(2)
)

odpowiedzi_2 = [
    'A. Pies dawa\u0142 mu do zrozumienia, \u017ce jest wa\u017cniejszy od oficer\u00f3w.',
    'B. Hektor jako jedyny na Placu traktowa\u0142 Nemeczka jak r\u00f3wnego sobie.',
    'C. Obaj \u2014 pies i Nemeczek \u2014 nie mieli \u017cadnego stopnia oficerskiego.',
    'D. Hektor pilnowa\u0142 twierdzy razem z Nemeczkiem podczas ka\u017cdego zebrania.'
]
for ans in odpowiedzi_2:
    add_paragraph(ans, space_after=Pt(2))

add_paragraph('', space_after=Pt(4))

# === ZADANIE 3 (FRAZEOLOGIZM) ===
add_mixed_paragraph([
    ('Zadanie 3. ', True, False),
    ('W fragmencie pojawia si\u0119 wyra\u017cenie: \u201eserce wali\u0142o mu jak m\u0142otem\u201d. Wyja\u015bnij znaczenie tego frazeologizmu i wska\u017c, jakie uczucia Nemeczka oddaje on w tym momencie.', False, False)
], space_after=Pt(4))
add_paragraph(LINE, space_after=Pt(2))

add_paragraph('', space_after=Pt(4))

# === ZADANIE 4 (SKLADNIA) ===
add_mixed_paragraph([
    ('Zadanie 4. ', True, False),
    ('Wypisz podmiot i orzeczenie z pierwszego zdania sk\u0142adowego poni\u017cszego zdania z\u0142o\u017conego.', False, False)
], space_after=Pt(2))

add_mixed_paragraph([
    ('\u201eSerce zabi\u0142o Nemeczkowi mocniej i zawaha\u0142 si\u0119, czy nie wr\u00f3ci\u0107 na ziemi\u0119.\u201d', False, True)
], space_after=Pt(4))

add_paragraph(
    'Podmiot: ' + '_' * 40 + '  Orzeczenie: ' + '_' * 40,
    space_after=Pt(2)
)

add_paragraph('', space_after=Pt(4))

# === ZADANIE 5 (NARRACJA) ===
add_mixed_paragraph([
    ('Zadanie 5. ', True, False),
    ('Na podstawie fragmentu okre\u015bl, jaki typ narratora prowadzi opowie\u015b\u0107 w ', False, False),
    ('Ch\u0142opcach z Placu Broni', False, True),
    ('. Nast\u0119pnie wypisz z tekstu jedno zdanie, kt\u00f3re dowodzi, \u017ce narrator ma dost\u0119p do wewn\u0119trznych prze\u017cy\u0107 bohatera.', False, False)
], space_after=Pt(4))

add_paragraph('Typ narratora: ' + '_' * 70, space_after=Pt(2))
add_paragraph('Zdanie z tekstu: ' + '_' * 67, space_after=Pt(2))
add_paragraph(LINE, space_after=Pt(2))

add_paragraph('', space_after=Pt(4))

# === ZADANIE 6 (PP_narrator) ===
add_mixed_paragraph([
    ('Zadanie 6. ', True, False),
    ('Okre\u015bl rodzaj narratora w ', False, False),
    ('Ch\u0142opcach z Placu Broni', False, True),
    ('. Wyja\u015bnij, czym charakteryzuje si\u0119 ten rodzaj narratora, i podaj jeden przyk\u0142ad z fragmentu, kt\u00f3ry potwierdza Tw\u00f3j wyb\u00f3r.', False, False)
], space_after=Pt(4))

add_paragraph('Rodzaj narratora: ' + '_' * 65, space_after=Pt(2))
add_paragraph('Charakterystyka: ' + '_' * 66, space_after=Pt(2))
add_paragraph('Przyk\u0142ad z tekstu: ' + '_' * 64, space_after=Pt(2))

add_paragraph('', space_after=Pt(4))

# === ZADANIE 7 (POWIAZANIE) ===
add_mixed_paragraph([
    ('Zadanie 7. ', True, False),
    ('Chor\u0105giewka Zwi\u0105zku odgrywa wa\u017cn\u0105 rol\u0119 w ca\u0142ej powie\u015bci. Opisz, w jaki spos\u00f3b motyw chor\u0105giewki lub walki o Plac Broni \u0142\u0105czy si\u0119 z postaw\u0105 innych postaci i z rozwojem wydarze\u0144 w ca\u0142ym utworze. Odwo\u0142aj si\u0119 do co najmniej dw\u00f3ch sytuacji z ca\u0142o\u015bci powie\u015bci \u2014 innych ni\u017c opisany fragment.', False, False)
], space_after=Pt(4))

add_paragraph(LINE, space_after=Pt(2))
add_paragraph(LINE, space_after=Pt(2))
add_paragraph(LINE, space_after=Pt(2))


# === ZAPIS ===
output_path = r"D:\Dydaktyka\ArkuszeAI\arkusz_chlopcy_08_rozdzII_Nemeczek_twierdza.docx"
doc.save(output_path)
print(f"Arkusz zapisany: {output_path}")
