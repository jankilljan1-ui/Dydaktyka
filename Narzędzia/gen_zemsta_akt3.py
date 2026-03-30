# -*- coding: utf-8 -*-
"""Generator arkusza E8 — Zemsta, Akt III scena I (id=7)"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT = r"D:\Dydaktyka\ArkuszeAI\Zemsta_AktIII_Sc1_Rejent_mularze.docx"
LINE = '_' * 90

doc = Document()

# ── Page setup ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

# ── Default style ───────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
font.color.rgb = RGBColor(0, 0, 0)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.space_before = Pt(0)

# ── Helper functions ────────────────────────────────────────

def tnr_run(paragraph, text, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return run

def add_p(text, bold=False, italic=False, alignment=None, space_after=Pt(2)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    if alignment:
        p.alignment = alignment
    tnr_run(p, text, bold, italic)
    return p

def add_drama(text, bold=False, italic=False, alignment=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    if alignment:
        p.alignment = alignment
    tnr_run(p, text, bold, italic)
    return p

def add_drama_mixed(parts, alignment=None):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    if alignment:
        p.alignment = alignment
    for text, b, i in parts:
        tnr_run(p, text, b, i)
    return p

def char_name(name):
    add_drama(name, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

def did(text):
    add_drama(text, italic=True)

def v(text):
    add_drama(text)

def empty():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

def answer_line():
    add_p(LINE)

# ── HEADER ──────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.space_before = Pt(0)
tnr_run(p, 'Aleksander Fredro, ', bold=True)
tnr_run(p, 'Zemsta', bold=True, italic=True)
tnr_run(p, ', akt III, scena I', bold=True)

# ── FRAGMENT ────────────────────────────────────────────────

add_drama('AKT TRZECI', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
empty()
add_drama('SCENA PIERWSZA', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
empty()

did('/ Rejent siedzi przy stoliku i pisze. Dwóch mularzy przy drzwiach stoi. /')
empty()
did('/ Rejent, MULARZE /')
empty()

char_name('REJENT')
v('Mój majstruniu, mówcie śmiało,')
v('Opiszemy sprawę całą;')
v('Na te ciężkie nasze czasy')
v('Boskim darem takie basy.')
v('Każdy kułak spieniężymy:')
v('Że was bito, wszyscy wiémy.')
empty()

char_name('MULARZ')
v('Niekoniecznie.')
empty()

char_name('REJENT')
v('                        Bili przecie,')
v('Mój majstruniu.')
empty()

char_name('MULARZ')
v('                        Niewyraźnie.')
empty()

char_name('REJENT')
v('Czegóż jeszcze wam nie stało?')
v('Bo machano dosyć raźnie.')
empty()

char_name('MULARZ')
v('Ot, szturknięto tam coś mało.')
empty()

char_name('DRUGI MULARZ')
v('Któż tam za to skarżyć zechce!')
empty()

char_name('REJENT')
v('Lecz kto szturka, ten nie łechce?')
empty()

char_name('MULARZ')
v('Ha! Zapewne.')
empty()

char_name('REJENT')
v('                        A więc bije?')
empty()

char_name('MULARZ')
v('Oczywiście.')
empty()

char_name('REJENT')
v('                        Komu kije')
v('Porachują kości w grzbiecie,')
v('Ten jest bity \u2014 wszak to wiecie?')
v('A kto bity, ten jest zbity?')
v('Co?')
empty()

char_name('MULARZ')
v('                        Ha! dobrze pan powiada,')
v('Ten jest zbity.')
empty()

char_name('REJENT')
v('                        Więc was zbili,')
v('To rzecz jasna, moi mili.')
empty()

char_name('MULARZ')
v('Ta, już jakoś tak wypada.')
empty()

char_name('REJENT')
did('/ napisawszy /')
empty()
v('Skaleczyli?')
empty()

char_name('MULARZ')
v('                        A, broń Boże!')
empty()

char_name('REJENT')
v('Nie, serdeńko?')
empty()

char_name('MULARZ')
v('                        Ach, nie.')
empty()

char_name('REJENT')
v('                        Przecie')
v('Znak, drapnięcie?')
empty()

char_name('MULARZ')
did('/ pomówiwszy z drugim /')
empty()
v('                        Znajdziem może.')
empty()

char_name('REJENT')
v('A drapnięcie, pewnie wiecie,')
v('Mała ranka, nic innego.')
empty()

char_name('MULARZ')
v('Tać, tak niby.')
empty()

char_name('REJENT')
v('                        Mała, wielka,')
v('Jednym słowem, rana wszelka')
v('Skąd pochodzi?')
empty()

char_name('MULARZ')
v('                        Niby\u2026 z tego\u2026')
empty()

char_name('REJENT')
v('Z skaleczenia.')
empty()

char_name('MULARZ')
v('                        Nie inaczej.')
empty()

char_name('REJENT')
v('Mieć więc ranę tyle znaczy,')
v('Co mieć ciało skaleczone:')
v('Że zaś raną jest drapnięcie,')
v('Więc zapewnić możem święcie,')
v('Że jesteście skaleczeni,')
v('Przez to chleba pozbawieni.')
empty()

char_name('MULARZ')
v('O! to znowu\u2026')
empty()

char_name('REJENT')
v('                        Pozbawiony')
v('Jesteś, bratku, i z przyczyny,')
v('Że ci nie dam okruszyny \u2014')
empty()
did('/ pisze /')
empty()
v('Zatem, zatem skaleczeni,')
v('Przez to chleba pozbawieni,')
v('Z matką \u2014 żoną \u2014 czworgiem dzieci.')
empty()

char_name('MULARZ')
v('Nie mam dzieci.')
empty()

char_name('DRUGI MULARZ')
v('                        Nie mam żony.')
empty()

char_name('REJENT')
v('Co? nie macie? \u2014 nic nie szkodzi \u2014')
v('Mieć możecie \u2014 tacy młodzi.')
empty()

char_name('MULARZ')
v('Ha!')
empty()

char_name('DRUGI MULARZ')
v('                        Tać prawda.')
empty()

char_name('REJENT')
did('/ napisawszy /')
empty()
v('                        Akt skończony.')
v('Teraz jeszcze zaświadczycie,')
v('Że nastawał na me życie.')
v('Stary Cześnik, jęty szałem,')
v('Strzelał do mnie.')
empty()

char_name('MULARZ')
v('                        Nie widziałem.')
empty()

char_name('REJENT')
v('Wołał strzelby.')
empty()

char_name('DRUGI MULARZ')
v('                        Nie słyszałem.')
empty()

char_name('MULARZ')
v('Wołał wprawdzie: \u00abDaj gwintówki!\u00bb \u2014')
v('Lecz chciał strzelać do makówki.')
empty()

char_name('REJENT')
v('Do makówki\u2026 do makówki\u2026')
v('No, no, dosyć tego będzie \u2014')
v('Świadków na to znajdę wszędzie \u2014')
v('Nie brak świadków na tym świecie,')
v('Teraz chodźcie \u2014 bliżej! bliżej! \u2014')
v('Znakiem krzyża podpiszecie. \u2014')
v('Michał Kafar trochę niżej \u2014')
v('Tak, tak \u2014 Maciej Miętus \u2014 pięknie! \u2014')
v('Za ten krzyżyk będą grosze,')
v('A Cześniczek z żółci pęknie.')
empty()

char_name('MULARZ')
v('Najpokorniej teraz proszę,')
v('Coś z dawnego nam przypadnie.')
empty()

char_name('REJENT')
v('Cześnik wszystko będzie płacił.')
empty()

char_name('MULARZ')
v('Jakoś, panie, to nieładnie\u2026')
empty()

char_name('REJENT')
v('Byleś wasze nic nie stracił.')
empty()

char_name('MULARZ')
v('Tum pracował\u2026')
empty()

char_name('REJENT')
did('/ popychając ich ku drzwiom /')
empty()
v('                        Idźże z Bogiem,')
v('Bo się poznasz z moim progiem.')
empty()

char_name('MULARZ')
v('Tu zapłata, każdy powie\u2026')
empty()

char_name('REJENT')
did('/ popychając ku drzwiom /')
empty()
v('Idź, serdeńko, bo cię trzepnę.')
empty()

char_name('MULARZ')
did('/ we drzwiach /')
empty()
v('Ależ przecie\u2026')
empty()

char_name('REJENT')
did('/ zamykając drzwi /')
empty()
v('                        Bądźcie zdrowi!')
v('Dobrzy ludzie, bądźcie zdrowi!')
empty()
did('/ wracając /')
empty()
v('Czapkę przedam, pas zastawię,')
v('A Cześnika stąd wykurzę;')
v('Będzie potem o tej sprawie')
v('Na wołowej pisał skórze.')
v('Lecz tajemne moje wieści,')
v('Jeśli wszystkie z prawdą zgodne,')
v('Tym, czym teraz serce pieści,')
v('Najboleśniej go ubodnę.')

# ── Separator before tasks ──────────────────────────────────
empty()
empty()

# ── ZADANIE 1 ───────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
tnr_run(p, 'Zadanie 1.', bold=True)
tnr_run(p, ' Oceń, które zdania opisują wydarzenia z powyższego fragmentu. Zaznacz ')
tnr_run(p, 'P', bold=True)
tnr_run(p, ' (prawda) lub ')
tnr_run(p, 'F', bold=True)
tnr_run(p, ' (fałsz).')

# P/F Table
table = doc.add_table(rows=3, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
for i, text in enumerate(['', 'P', 'F']):
    cell = table.rows[0].cells[i]
    pr = cell.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tnr_run(pr, text, bold=True)

statements = [
    'Mularze zgodnie przyznali, że zostali poważnie pobici i skaleczeni przez hajduków Cześnika.',
    'Nie znalazłszy naocznych świadków strzelania, Rejent oznajmił, że wyszuka ich sam.'
]

for i, stmt in enumerate(statements):
    cell = table.rows[i+1].cells[0]
    pr = cell.paragraphs[0]
    tnr_run(pr, f'{i+1}. {stmt}')
    for j in [1, 2]:
        cell = table.rows[i+1].cells[j]
        pr = cell.paragraphs[0]
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Set column widths
for row in table.rows:
    row.cells[0].width = Cm(15.0)
    row.cells[1].width = Cm(1.5)
    row.cells[2].width = Cm(1.5)

# Table borders
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_table_borders(tbl):
    tbl_pr = tbl._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    tbl_pr.append(borders)

set_table_borders(table)

empty()

# ── ZADANIE 2 ───────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
tnr_run(p, 'Zadanie 2.', bold=True)
tnr_run(p, ' W jakim celu Rejent sporządza pismo podczas rozmowy z mularzami? Zaznacz właściwą odpowiedź.')

options = [
    'A) Aby nagrodzić mularzy za pracę przy remoncie muru granicznego.',
    'B) Aby przygotować doniesienie sądowe skierowane przeciwko Cześnikowi.',
    'C) Aby zawrzeć ugodę z Cześnikiem w sprawie odszkodowania dla mularzy.',
    'D) Aby wyrównać zaległe należności mularzom zgodnie z umową.'
]
for opt in options:
    add_p(opt)

empty()

# ── ZADANIE 3 ───────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
tnr_run(p, 'Zadanie 3.', bold=True)
tnr_run(p, ' W ostatnich słowach Rejenta pojawia się wyrażenie: ')
tnr_run(p, '\u201eBędzie potem o tej sprawie / Na wołowej pisał skórze\u201d', italic=True)
tnr_run(p, '. Wyjaśnij znaczenie użytego frazeologizmu. Jaki stosunek Rejenta do Cześnika oddaje użycie tego wyrażenia?')

add_p('Znaczenie:')
answer_line()
add_p('Stosunek Rejenta do Cześnika:')
answer_line()

empty()

# ── ZADANIE 4 ───────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
tnr_run(p, 'Zadanie 4.', bold=True)
tnr_run(p, ' Przeczytaj uważnie podane zdanie, a następnie wskaż zdanie podrzędne. Określ jego rodzaj i wyjaśnij, jaką relację wyraża ono wobec zdania nadrzędnego.')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
tnr_run(p, 'Idź, serdeńko, bo cię trzepnę.', italic=True)

add_p('Zdanie podrzędne:')
answer_line()
add_p('Rodzaj zdania podrzędnego:')
answer_line()
add_p('Relacja wobec zdania nadrzędnego:')
answer_line()

empty()

# ── ZADANIE 5 ───────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
tnr_run(p, 'Zadanie 5.', bold=True)
tnr_run(p, ' Spośród lektur obowiązkowych \u2014 innych niż ')
tnr_run(p, 'Zemsta', italic=True)
tnr_run(p, ' \u2014 wybierz tę, w której jeden z bohaterów posługuje się podstępem lub kłamstwem, aby osiągnąć swój cel. Podaj tytuł lektury i imię bohatera. Uzasadnij wybór, przywołując konkretną sytuację z tekstu.')

add_p('Tytuł lektury i imię bohatera:')
answer_line()
add_p('Uzasadnienie:')
answer_line()
answer_line()
answer_line()

empty()

# ── ZADANIE 6 ───────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
tnr_run(p, 'Zadanie 6.', bold=True)
tnr_run(p, ' W tekście dramatycznym, obok dialogów i monologów, obecne są inne charakterystyczne elementy. Podaj nazwę tych elementów, które w powyższym fragmencie ')
tnr_run(p, 'Zemsty', italic=True)
tnr_run(p, ' pojawiają się w nawiasach ukośnych. Wyjaśnij, jaką funkcję pełnią w dramacie. Przytocz dwa przykłady z tekstu.')

add_p('Nazwa:')
answer_line()
add_p('Funkcja:')
answer_line()
add_p('Przykład 1:')
answer_line()
add_p('Przykład 2:')
answer_line()

empty()

# ── ZADANIE 7 ───────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
tnr_run(p, 'Zadanie 7.', bold=True)
tnr_run(p, ' Ponumeruj (1\u20134) poniższe wydarzenia z ')
tnr_run(p, 'Zemsty', italic=True)
tnr_run(p, ' według kolejności, w jakiej pojawiają się w dramacie. Wśród wydarzeń nie ma zdarzeń z powyższego fragmentu.')

events = [
    '___ Papkin dostarcza Cześnikowi jako \u201ejeńca\u201d komisarza Rejenta \u2014 Wacława.',
    '___ Cześnik wydaje Papkinowi polecenie, by spędził mularzy z muru granicznego.',
    '___ Podstolina ogłasza, że odrzuca starania Cześnika i przyjmuje oświadczyny Rejenta.',
    '___ Cześnik i Rejent dochodzą do zgody i postanawiają, że Wacław weźmie ślub z Klarą.'
]
for ev in events:
    add_p(ev)

# ── SAVE ────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Zapisano: {OUTPUT}")
