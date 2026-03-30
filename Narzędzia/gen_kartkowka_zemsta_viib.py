# -*- coding: utf-8 -*-
"""
Generator Kartkowki — Zemsta (Aleksander Fredro)
Klasa: VIIB | Zakres: cala lektura | 3 pytania problemowe, 12 pkt
Data: 2026-03-28
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"D:\Dydaktyka\ArkuszeAI\Zemsta_Kartkowka_VIIB.docx"

doc = Document()

# Marginesy 1.27 cm
for section in doc.sections:
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(2)

FONT = 'Times New Roman'
SIZE = Pt(12)
DOTS = '............................................................................................................'


def p(text='', bold=False, italic=False, align=None, sa=Pt(2), sb=Pt(0), size=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = FONT
    run.font.size = size if size else SIZE
    run.bold = bold
    run.italic = italic
    if align:
        para.alignment = align
    para.paragraph_format.space_after = sa
    para.paragraph_format.space_before = sb
    return para


def mixed(parts, align=None, sa=Pt(2), sb=Pt(0)):
    para = doc.add_paragraph()
    for text, b, i in parts:
        run = para.add_run(text)
        run.font.name = FONT
        run.font.size = SIZE
        run.bold = b
        run.italic = i
    if align:
        para.alignment = align
    para.paragraph_format.space_after = sa
    para.paragraph_format.space_before = sb
    return para


def dots(n, first_sb=Pt(6)):
    for i in range(n):
        para = doc.add_paragraph()
        run = para.add_run(DOTS)
        run.font.name = FONT
        run.font.size = SIZE
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.space_before = first_sb if i == 0 else Pt(0)


def add_bottom_border(para):
    """Dodaje linie pod akapitem (separator)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def page_break():
    doc.add_page_break()


def add_grades_table():
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    data = [
        ('Ocena', 'Punkty'),
        ('Celuj\u0105cy', '12'),
        ('Bardzo dobry', '10\u201311'),
        ('Dobry', '8\u20139'),
        ('Dostateczny', '6\u20137'),
        ('Dopuszczaj\u0105cy', '4\u20135'),
        ('Niedostateczny', '0\u20133'),
    ]
    for i, (a, b) in enumerate(data):
        row = table.rows[i].cells
        row[0].text = a
        row[1].text = b
        for cell in row:
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell.text)
            run.font.name = FONT
            run.font.size = SIZE
            if i == 0:
                run.bold = True
    from docx.shared import Cm as C
    for row in table.rows:
        row.cells[0].width = C(4.0)
        row.cells[1].width = C(2.0)


# ══════════════════════════════════════════════════
# STRONA 1 — Kartkowka
# ══════════════════════════════════════════════════

# Dane ucznia
mixed([
    ('Imi\u0119 i nazwisko: ..........................................   ', False, False),
    ('Klasa: ............', False, False),
], sa=Pt(4))
p('Data: .........................', sa=Pt(10))

# Tytul lektury i punktacja
sep_para = p('\u017bzemsta, Aleksander Fredro \u2014 zakres: ca\u0142a lektura', italic=True, sa=Pt(2))
add_bottom_border(sep_para)
p('Punktacja: 0\u201312 pkt\u2003|\u2003Czas: 15 min', italic=True, sa=Pt(12))

# ── Pytanie 1 ──
mixed([('Pytanie 1.', True, False), (' (0\u20134 pkt)', False, False)], sb=Pt(4), sa=Pt(4))
p('Dlaczego Cze\u015bnik Raptusiewicz nie zgadza si\u0119 na napraw\u0119 muru przez Rejenta Milczka, cho\u0107 mur jest zrujnowany i wymaga remontu? Co tak naprawd\u0119 kryje si\u0119 za jego sprzeciwem? Odwo\u0142aj si\u0119 do wydarze\u0144 z dramatu.', sa=Pt(4))
dots(4)

# ── Pytanie 2 ──
mixed([('Pytanie 2.', True, False), (' (0\u20134 pkt)', False, False)], sb=Pt(8), sa=Pt(4))
p('W jaki spos\u00f3b charakter Papkina \u2014 jego pr\u00f3\u017cno\u015b\u0107 i tch\u00f3rzostwo \u2014 wp\u0142ywa na rozw\u00f3j konfliktu mi\u0119dzy Cze\u015bnikiem a Rejentem? Podaj co najmniej jeden przyk\u0142ad sytuacji, w kt\u00f3rej Papkin zamiast pom\u00f3c, komplikuje spraw\u0119.', sa=Pt(2))
p('Odwo\u0142aj si\u0119 do konkretnego wydarzenia z lektury.', italic=True, sa=Pt(4))
dots(5)

# ── Pytanie 3 ──
mixed([('Pytanie 3.', True, False), (' (0\u20134 pkt)', False, False)], sb=Pt(8), sa=Pt(4))
p('Czy Rejent Milczek jest bardziej przebieg\u0142ym i niebezpiecznym przeciwnikiem ni\u017c Cze\u015bnik Raptusiewicz? Zajmij stanowisko i uzasadnij je, odwo\u0142uj\u0105c si\u0119 do konkretnych zachowa\u0144 obu postaci w \u017bZem\u015bcie\u201d Aleksandra Fredry.', sa=Pt(2))
p('Odwo\u0142aj si\u0119 do konkretnego wydarzenia z lektury.', italic=True, sa=Pt(4))
dots(6)


# ══════════════════════════════════════════════════
# STRONA 2 — Klucz odpowiedzi
# ══════════════════════════════════════════════════

page_break()

klucz = p('KLUCZ ODPOWIEDZI \u2014 KARTKÓWKA', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(13), sa=Pt(2))
add_bottom_border(klucz)
p('(tylko dla nauczyciela)', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(10))

# Pytanie 1 — klucz
p('Pytanie 1.', bold=True, sb=Pt(4), sa=Pt(2))
p('Odpowied\u017a wzorcowa: Cze\u015bnik sprzeciwia si\u0119 naprawie muru, bo traktuje to jako ust\u0119pstwo wobec znienawidzonego s\u0105siada. Prawdziw\u0105 przyczyn\u0105 nie jest troska o mur, lecz duma i ch\u0119\u0107 dominacji \u2014 zgoda oznacza\u0142aby przyznanie racji Rejentowi i utrat\u0119 twarzy.', sa=Pt(2))
p('2 pkt: ucze\u0144 wskazuje honor/dum\u0119 jako motywacj\u0119 + odwo\u0142uje si\u0119 do konkretnej sceny.', sa=Pt(1))
p('1 pkt: og\u00f3lne wskazanie na konflikt s\u0105siedzki bez uzasadnienia motywacji.', sa=Pt(1))
p('0 pkt: odpowied\u017a odtw\u00f3rcza lub brak odwo\u0142ania do tekstu.', sa=Pt(1))
p('Typowe b\u0142\u0119dy: \u201ebo nie lubi\u0142 Rejenta\u201d (zbyt og\u00f3lne); mylenie powodu formalnego z motywacj\u0105 psychologiczn\u0105.', italic=True, sa=Pt(8))

# Pytanie 2 — klucz
p('Pytanie 2.', bold=True, sb=Pt(4), sa=Pt(2))
p('Odpowied\u017a wzorcowa: Papkin, wys\u0142any przez Cze\u015bnika jako pose\u0142 do Rejenta, zamiast negocjowa\u0107 \u2014 chwali si\u0119 fa\u0142szywymi zas\u0142ugami i daje si\u0119 rozbroi\u0107 s\u0142owem. Jego brawura jest pozorowana, co sprawia, \u017ce misja dyplomatyczna zamienia si\u0119 w komiczn\u0105 pora\u017ck\u0119 i zaostrza konflikt.', sa=Pt(2))
p('2 pkt: uczeń nazywa cech\u0119 (pr\u00f3\u017cno\u015b\u0107/tch\u00f3rzostwo) + podaje konkretn\u0105 scen\u0119 (np. wizyta u Rejenta, list z wyzwaniem).', sa=Pt(1))
p('1 pkt: wymienia cech\u0119 bez przyk\u0142adu lub przyk\u0142ad bez analizy jej wp\u0142ywu.', sa=Pt(1))
p('0 pkt: streszczenie bez analizy charakteru.', sa=Pt(1))
p('Typowe b\u0142\u0119dy: skupienie na tym, co Papkin m\u00f3wi o sobie, bez pokazania, jak to wp\u0142ywa na fabu\u0142\u0119.', italic=True, sa=Pt(8))

# Pytanie 3 — klucz
p('Pytanie 3.', bold=True, sb=Pt(4), sa=Pt(2))
p('Odpowied\u017a wzorcowa: Rejent jest gro\u017aniejszym przeciwnikiem, bo dzia\u0142a skrycie i metodycznie \u2014 manipuluje sytuacj\u0105, planuje przej\u0119cie zamku, ukrywa prawdziwe zamiary za pozorn\u0105 grzeczno\u015bci\u0105. Cze\u015bnik jest gwa\u0142towny, lecz otwarty; Rejent jest zimny i wyrachowany, co czyni go bardziej niebezpiecznym.', sa=Pt(2))
p('2 pkt: teza + por\u00f3wnanie zachowa\u0144 obu postaci z odwo\u0142aniem do tekstu.', sa=Pt(1))
p('1 pkt: teza bez por\u00f3wnania lub por\u00f3wnanie bez odwo\u0142ania do tekstu.', sa=Pt(1))
p('0 pkt: brak stanowiska lub streszczenie bez oceny.', sa=Pt(1))
p('Typowe b\u0142\u0119dy: uznanie obu za jednakowo winnych bez argumentacji; pomini\u0119cie konkretnych dzia\u0142a\u0144 Rejenta.', italic=True, sa=Pt(10))

# Tabela ocen
p('Progi ocen:', bold=True, sa=Pt(4))
add_grades_table()

doc.save(OUTPUT_PATH)
print(f'Zapisano: {OUTPUT_PATH}')
