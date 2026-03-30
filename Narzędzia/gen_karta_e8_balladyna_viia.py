# -*- coding: utf-8 -*-
"""
Generator Karty Pracy E8 — Balladyna (Juliusz Słowacki)
Fragment: scena zabójstwa Aliny (akt II)
Klasa: VIIA
5 kategorii analitycznych, 10 pkt
Data: 2026-03-28
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_PATH = r"D:\Dydaktyka\ArkuszeAI\Balladyna_ZabojstwoAliny_KartaE8_VIIA.docx"

doc = Document()

# Marginesy 1.27 cm
for section in doc.sections:
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

# Style domyslne
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(2)

FONT = 'Times New Roman'
SIZE = Pt(12)
DOTS = '............................................................................................................'


# ── Funkcje pomocnicze ──

def p(text='', bold=False, italic=False, align=None, sa=Pt(2), sb=Pt(0), indent_cm=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = FONT
    run.font.size = SIZE
    run.bold = bold
    run.italic = italic
    if align:
        para.alignment = align
    para.paragraph_format.space_after = sa
    para.paragraph_format.space_before = sb
    if indent_cm:
        para.paragraph_format.left_indent = Cm(indent_cm)
    return para


def dots(n=3, first_sb=Pt(6)):
    for i in range(n):
        sb = first_sb if i == 0 else Pt(0)
        para = doc.add_paragraph()
        run = para.add_run(DOTS)
        run.font.name = FONT
        run.font.size = SIZE
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.space_before = sb


def page_break():
    doc.add_page_break()


def zadanie_header(nr, kategoria, punkty):
    p(f'Zadanie {nr}. {kategoria}', bold=True, sb=Pt(10), sa=Pt(2))
    p(f'(0\u2013{punkty} pkt)', sa=Pt(4))


def add_table_pf(rows_data):
    """rows_data = list of (stwierdzenie,)"""
    table = doc.add_table(rows=len(rows_data) + 1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for cell, txt in zip(hdr, ['Stwierdzenie', 'P', 'F']):
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = FONT
        cell.paragraphs[0].runs[0].font.size = SIZE
    for i, (stmt,) in enumerate(rows_data):
        row = table.rows[i + 1].cells
        row[0].text = stmt
        row[1].text = ''
        row[2].text = ''
        for cell in row:
            for run in cell.paragraphs[0].runs:
                run.font.name = FONT
                run.font.size = SIZE
    widths = [Cm(11.5), Cm(1.0), Cm(1.0)]
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w
    # odstep po tabeli
    p('', sa=Pt(4))


def add_klucz_table(rows_data):
    """rows_data = list of (nr, kategoria, odpowiedz, punkty)"""
    table = doc.add_table(rows=len(rows_data) + 1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for cell, txt in zip(hdr, ['Nr', 'Kategoria', 'Odpowiedź wzorcowa', 'Pkt']):
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = FONT
        cell.paragraphs[0].runs[0].font.size = SIZE
    for i, (nr, kat, odp, pkt) in enumerate(rows_data):
        row = table.rows[i + 1].cells
        for cell, val in zip(row, [str(nr), kat, odp, str(pkt)]):
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.name = FONT
                run.font.size = SIZE
    widths = [Cm(0.8), Cm(2.8), Cm(9.0), Cm(0.9)]
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w


# ══════════════════════════════════════════════════
# STRONA 1 — Fragment literacki
# ══════════════════════════════════════════════════

def drama_verse(postac=None, didaskalia=None, wersy=None):
    if postac:
        p(postac, bold=True, sa=Pt(1))
    if didaskalia:
        p(didaskalia, italic=True, sa=Pt(1), indent_cm=0.5)
    if wersy:
        for wers in wersy:
            p(wers, sa=Pt(1), indent_cm=0.5)


drama_verse(postac='BALLADYNA',
            didaskalia='/ z wsciekloscia natrętną /',
            wersy=['Oddaj mi ten dzbanek.'])

drama_verse(postac='ALINA',
            wersy=['Siostro?\u2026'])

drama_verse(postac='BALLADYNA',
            wersy=['Oddaj mi\u2026 bo!\u2026'])

drama_verse(postac='ALINA',
            didaskalia='/ z dziecinnym naigrawaniem się /',
            wersy=[
                'Bo!\u2026 i cóż będzie?..',
                'Bo?\u2026 Nie masz malin, więc suche żołędzie',
                'Uzbierasz w dzbanek \u2014 czy wierzbowe liście?\u2026',
                'I tak\u2026 ja prędzej biegam i przez miedzę',
                'Ubiegnę ciebie\u2026'
            ])

drama_verse(postac='BALLADYNA',
            wersy=['Ty?\u2026'])

drama_verse(postac='ALINA',
            wersy=[
                'A oczewiście,',
                'Że ciebie w locie, siostrzyczko, wyprzedzę\u2026'
            ])

drama_verse(postac='BALLADYNA',
            wersy=['Ty!'])

drama_verse(postac='ALINA',
            wersy=[
                'O! nie zbliżaj się do mnie z takiemi',
                'Oczyma\u2026 Nie wiem\u2026 ja się ciebie boję.'
            ])

drama_verse(postac='BALLADYNA',
            didaskalia='/ zbliża się i bierze ją za rękę /',
            wersy=[
                'I ja się boję\u2026 połóż się na ziemi\u2026',
                'Połóż\u2026 ha!'
            ])

p('/ zabija /', italic=True, sa=Pt(1), indent_cm=0.5)

drama_verse(postac='ALINA',
            wersy=['Puszczaj!\u2026 oh!\u2026 konam\u2026'])

p('/ pada /', italic=True, sa=Pt(1), indent_cm=0.5)

drama_verse(postac='BALLADYNA',
            wersy=['Co moje', 'Ręce zrobiły?\u2026 O!\u2026'])

drama_verse(postac='GŁOS Z WIERZBY',
            wersy=['Jezus Maryja\u2026'])

drama_verse(postac='BALLADYNA',
            didaskalia='/ przerażona /',
            wersy=[
                'Kto to?\u2026 zawołał ktoś?\u2026 czy to ja sama',
                'Za siebie samą modliłam się?\u2026 Żmija,',
                'Kobieta, siostra \u2014 nie siostra\u2026 Krwi plama',
                'Tu \u2014 i tu \u2014 i tu \u2014'
            ])

p('/ pokazując na czoło, plami je palcem /', italic=True, sa=Pt(1), indent_cm=0.5)

p('i tu. \u2014 Ktoż zabija', sa=Pt(1), indent_cm=0.5)

for wers in [
    'Za malin dzbanek siostrę?\u2026 Jeśli z bora',
    'Kto tak zapyta? powiem \u2014 ja. \u2014 Nie mogę',
    'Skłamać i powiem: ja! \u2014 Jak to ja?\u2026 Wczora',
    'Mogłabym przysiąc, że nie\u2026 W las!\u2026 w las!\u2026 w drogę,',
    'Wczorajsze serce niechaj się za ciebie',
    'Modli. \u2014 Ach jam się wczoraj nie modliła.',
    'To źle! źle! \u2014 dzisiaj już nie czas\u2026 Na niebie',
    'Jest Bóg\u2026 zapomnę, że jest, będę żyła,',
    'Jakby nie było Boga.',
]:
    p(wers, sa=Pt(1), indent_cm=0.5)

p('/ odbiega w las /', italic=True, sa=Pt(8), indent_cm=0.5)

p('Juliusz Słowacki, Balladyna, akt II',
  italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT, sa=Pt(2))


# ══════════════════════════════════════════════════
# STRONA 2 — Zadania 1–5
# ══════════════════════════════════════════════════

page_break()

# Zadanie 1. Kontekst — P/F (0–1 pkt)
zadanie_header(1, 'Kontekst', 1)
p('Oceń prawdziwość poniższych stwierdzeń dotyczących fragmentu. Zaznacz P (prawda) lub F (fałsz).', sa=Pt(6))
add_table_pf([
    ('Balladyna tłumaczy, że nóż wzięła do obrony przed wężami w malinach.',),
    ('Alina gotowa jest oddać maliny siostrze bez żadnych warunków.',),
    ('Po zabójstwie Balladyna postanawia uciec w las i zapomnieć o Bogu.',),
])

# Zadanie 2. Architektura tekstu — środek stylistyczny (0–2 pkt)
zadanie_header(2, 'Architektura tekstu', 2)
p('W podanym cytacie wskaż środek stylistyczny i wyjaśnij, jaką pełni funkcję w tym fragmencie.', sa=Pt(4))
p('\u201eKtoż zabija / Za malin dzbanek siostrę?\u201d', italic=True, indent_cm=1.0, sa=Pt(6))
p('Środek stylistyczny: ', sa=Pt(2))
dots(2)
p('Funkcja: ', sb=Pt(6), sa=Pt(2))
dots(2)

# Zadanie 3. Portret psychologiczny — motywacja (0–2 pkt)
zadanie_header(3, 'Portret psychologiczny', 2)
p('Na podstawie zachowania Balladyny w podanym fragmencie wyjaśnij, co motywuje ją do zabójstwa siostry. Odwołaj się do dwóch szczegółów z tekstu.', sa=Pt(6))
dots(4)

# Zadanie 4. Kod kulturowy — motyw literacki (0–2 pkt)
zadanie_header(4, 'Kod kulturowy', 2)
p('Jaki motyw literacki rozpoznajesz w wypowiedzi Balladyny bezpośrednio po dokonaniu zabójstwa? Nazwij go i wyjaśnij, w jaki sposób realizuje się w tym fragmencie.', sa=Pt(6))
dots(4)

# Zadanie 5. Argumentacja — krótka wypowiedź (0–3 pkt)
zadanie_header(5, 'Argumentacja', 3)
p('Czy sądzisz, że decyzja Balladyny, by \u201ezapomnieć, że jest Bóg\u201d, pozwoliła jej żyć bez wyrzutów sumienia? Uzasadnij swoje stanowisko, odwołując się do fragmentu oraz do dalszych losów bohaterki w całości dramatu Juliusza Słowackiego \u201eBalladyna\u201d.', sa=Pt(6))
dots(8)


# ══════════════════════════════════════════════════
# STRONA 3 — Klucz odpowiedzi
# ══════════════════════════════════════════════════

page_break()

p('KLUCZ ODPOWIEDZI', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(6))
p('(tylko dla nauczyciela)', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=Pt(12))

add_klucz_table([
    (1, 'Kontekst',
     'P \u2014 F \u2014 P\nStwierdzenie 1: Balladyna mówi „ten nóż to na węża w malinach" \u2192 P.\nStwierdzenie 2: Alina oddałaby maliny tylko po pocałunku \u2192 F.\nStwierdzenie 3: Balladyna ucieka i mówi, że zapomni o Bogu \u2192 P.',
     '0\u20131'),
    (2, 'Architektura tekstu',
     'Pytanie retoryczne. (1 pkt)\nFunkcja: wyraża wewnętrzny dialog Balladyny z własnym sumieniem; bohaterka jakby sama sobie stawia pytanie o sens zbrodni, co pokazuje jej narastający strach i poczucie winy. (1 pkt)',
     '0\u20132'),
    (3, 'Portret psychologiczny',
     'Motywacja: ambicja i zazdrość \u2014 Balladyna chce wygrać konkurs na maliny, by pojąć Kirkora za męża i awansować społecznie. (1 pkt)\nOdwołanie do tekstu: np. żądanie oddania dzbanka, narastająca wściekłość, kpiny Aliny o sukcesie w biegu. (1 pkt)',
     '0\u20132'),
    (4, 'Kod kulturowy',
     'Motyw winy i sumienia / motyw zbrodniarza ściganego przez wyrzuty sumienia (analogia do Makbeta Szekspira). (1 pkt)\nRealizacja: Balladyna widzi wszędzie „Krwi plamę", prowadzi gorączkowy monolog, świadomie postanawia wyrzec się Boga, by stłumić sumienie. (1 pkt)',
     '0\u20132'),
    (5, 'Argumentacja',
     'Teza (np. Nie \u2014 sumienie powraca wbrew jej woli). (1 pkt)\nArgument z fragmentu: słowa „To źle! źle!", plama krwi, strach przed głosem z wierzby. (1 pkt)\nOdwołanie do całości: sceny sądowe, publiczne wyznanie winy, wyrok piorunem \u2014 sumienie Balladyny nigdy nie milknie. (1 pkt)',
     '0\u20133'),
])

doc.save(OUTPUT_PATH)
print(f'Zapisano: {OUTPUT_PATH}')
